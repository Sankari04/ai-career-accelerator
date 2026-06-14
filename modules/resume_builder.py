"""
resume_builder.py
AI Resume Builder — generates a one-page ATS-optimised resume,
stores it in session state, and allows repeated PDF / DOCX downloads
without re-generation.
"""

import io
import streamlit as st
from fpdf import FPDF
from docx import Document
from utils.api_client import generate_response
from utils import session_state as ss


# ── PDF helper ────────────────────────────────────────────────────────────────

class _PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 15)
        self.cell(0, 10, "Professional Resume", 0, 1, "C")
        self.ln(3)


def _build_pdf(content: str) -> bytes:
    """Convert markdown-ish resume text to a clean PDF."""
    pdf = _PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=12)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        # Strip markdown bold markers
        line = line.replace("**", "")
        # Encode to latin-1 safely
        line = line.encode("latin-1", "replace").decode("latin-1")

        if line.startswith("# "):
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 8, line[2:].strip(), ln=1)
            pdf.set_font("Arial", size=10)
        elif line.startswith("## "):
            pdf.set_font("Arial", "B", 11)
            pdf.cell(0, 7, line[3:].strip(), ln=1)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.set_font("Arial", size=10)
        elif line.startswith("### "):
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, line[4:].strip(), ln=1)
            pdf.set_font("Arial", size=10)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.multi_cell(0, 5, f"  \u2022 {line[2:].strip()}")
        else:
            pdf.multi_cell(0, 5, line)

    return pdf.output(dest="S").encode("latin-1")


def _build_docx(content: str) -> bytes:
    """Convert markdown-ish resume text to a DOCX document."""
    doc = Document()
    doc.add_heading("Professional Resume", 0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        line = line.replace("**", "")

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ── Prompt builder ─────────────────────────────────────────────────────────────

def _build_prompt(name, email, phone, linkedin, education, skills,
                  experience, projects, certifications) -> str:
    return f"""
You are an expert resume writer and ATS specialist. Generate a strictly ONE-PAGE,
ATS-optimised resume in clean Markdown for the candidate below.

STRICT RULES — never break these:
- Total output must fit on one printed A4 page (≈ 550–600 words max).
- Professional Summary: exactly 3–4 sentences, no more.
- Core Competencies: list 8–12 comma-separated skills on one line.
- Work Experience bullet points: strong action verbs, quantify where possible,
  max 3–4 bullets per role.
- Projects: max 2–3 bullets each; focus on tech used and measurable outcome.
- Do NOT include filler phrases, generic objectives, or padding.
- Output ONLY the resume — no explanations, no preamble.

Candidate details:
Name: {name}
Email: {email} | Phone: {phone} | LinkedIn: {linkedin}
Education: {education}
Skills: {skills}
Experience: {experience}
Projects: {projects}
Certifications: {certifications}

Use this exact structure:

# {name}
{email} | {phone} | {linkedin}

## Professional Summary
(3–4 impactful sentences)

## Core Competencies
(comma-separated list)

## Professional Experience
### Job Title — Company | Dates
- bullet
- bullet
- bullet

## Key Projects
### Project Name | Tech Stack
- bullet
- bullet

## Education & Certifications
(formatted clearly)
"""


# ── Main render function ───────────────────────────────────────────────────────

def render():
    st.markdown("<div class='animate-fade-in'>", unsafe_allow_html=True)

    # Header
    c1, c2 = st.columns([3, 1])
    with c1:
        st.markdown("<h1 style='margin-bottom:0;'>📄 AI Resume Builder</h1>",
                    unsafe_allow_html=True)
        st.markdown("<p style='color:#94a3b8;font-size:1.1rem;'>"
                    "Craft a one-page, ATS-optimised resume — stored so you "
                    "can download anytime without re-generating.</p>",
                    unsafe_allow_html=True)
    with c2:
        st.markdown("<div style='text-align:right;padding-top:20px;'>"
                    "<span class='premium-badge'>PDF & DOCX Export</span></div>",
                    unsafe_allow_html=True)

    # ── Input form ────────────────────────────────────────────────────────────
    with st.form("resume_form"):
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<h3 style='margin-top:0;color:#a78bfa;'>1. Personal Details</h3>",
                    unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            name  = st.text_input("Full Name *", placeholder="Jane Doe")
            email = st.text_input("Email *",     placeholder="jane@example.com")
        with c2:
            phone    = st.text_input("Phone",    placeholder="+1 234 567 890")
            linkedin = st.text_input("LinkedIn", placeholder="linkedin.com/in/janedoe")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<h3 style='margin-top:0;color:#a78bfa;'>2. Education & Skills</h3>",
                    unsafe_allow_html=True)
        c3, c4 = st.columns(2)
        with c3:
            education = st.text_area("Education",
                placeholder="B.S. Computer Science, MIT, 2020–2024", height=110)
        with c4:
            skills = st.text_area("Core Skills (comma-separated) *",
                placeholder="Python, React, SQL, Docker, AWS", height=110)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<h3 style='margin-top:0;color:#a78bfa;'>3. Experience & Projects</h3>",
                    unsafe_allow_html=True)
        experience     = st.text_area("Work Experience",
            placeholder="Software Engineer @ TechCorp\n- Built X, improved Y by 20%",
            height=140)
        projects       = st.text_area("Key Projects",
            placeholder="AI Resume Builder — Python, Streamlit\n- Reduced ...",
            height=110)
        certifications = st.text_input("Certifications / Awards",
            placeholder="AWS Certified Solutions Architect")
        st.markdown("</div>", unsafe_allow_html=True)

        submitted = st.form_submit_button(
            "✨ Generate One-Page Resume", use_container_width=True)

    # ── Generation ────────────────────────────────────────────────────────────
    if submitted:
        if not name or not email or not skills:
            st.error("⚠️ Name, Email, and Skills are required.")
        else:
            with st.spinner("🧠 Crafting your one-page resume…"):
                prompt   = _build_prompt(name, email, phone, linkedin,
                                         education, skills, experience,
                                         projects, certifications)
                content  = generate_response(prompt)
                pdf_bytes  = _build_pdf(content)
                docx_bytes = _build_docx(content)

            # Persist everything in session state
            ss.set_value("resume", "content",  content)
            ss.set_value("resume", "pdf",      pdf_bytes)
            ss.set_value("resume", "docx",     docx_bytes)
            ss.set_value("resume", "filename", name.replace(" ", "_"))

            st.success("✅ Resume generated and saved to your session!")

    # ── Display stored result (persists across navigation) ────────────────────
    _show_resume_result()
    st.markdown("</div>", unsafe_allow_html=True)


def _show_resume_result():
    """Render the preview and download buttons from session state."""
    data = ss.get("resume")
    if not data.get("content"):
        return

    st.markdown("<div class='glass-card animate-fade-in'>", unsafe_allow_html=True)
    st.markdown("### 👁️ Resume Preview")
    st.markdown(
        "<div style='background:rgba(15,23,42,0.8);padding:28px;"
        "border-radius:12px;border:1px solid rgba(255,255,255,0.1);'>",
        unsafe_allow_html=True)
    st.markdown(data["content"])
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='glass-card animate-fade-in'>", unsafe_allow_html=True)
    st.markdown("### 📥 Download Your Resume")
    st.markdown("<p style='color:#94a3b8;margin-bottom:16px;'>"
                "Both formats are generated once and stay available for "
                "the entire session.</p>",
                unsafe_allow_html=True)

    fname = data.get("filename", "resume")
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "📄 Download PDF",
            data=data["pdf"],
            file_name=f"{fname}_Resume.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="resume_pdf_dl",
        )
    with c2:
        st.download_button(
            "📝 Download DOCX",
            data=data["docx"],
            file_name=f"{fname}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
            use_container_width=True,
            key="resume_docx_dl",
        )
    st.markdown("</div>", unsafe_allow_html=True)
